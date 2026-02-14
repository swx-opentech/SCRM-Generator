from docx import *
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from docx.document import Document as DocumentObject
from file_manager import File_Manager

class DocumentWriter:
    software_name = "[软件名称]"
    software_version = "[软件版本]"
    current_codeblock = 0
    total_codeblock = 0

    # 创建文档
    @classmethod
    def Create(cls) -> DocumentObject:
        with open("source-code_model.docx", "rb") as f:
            file_stream = io.BytesIO(f.read())
        doc = Document(file_stream)

        if len(doc.paragraphs) > 0 and doc.paragraphs[0].text.strip() == "":
            p_element = doc.paragraphs[0]._element
            p_element.getparent().remove(p_element)
        return doc
    

    @classmethod
    def __Add_Title__(cls, doc: DocumentObject, title_text: str, title_order: int) -> None:
        # 处理标题
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        title_run = title_para.add_run(f"第{title_order}部分 {title_text}")
        title_run.font.size = Pt(18)
        title_run.font.name = "Consolas"
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
        title_run.font.bold = True



    # header_order 就是 section 的顺序，从 1 开始。注意：section = header_order - 1
    @classmethod
    def __Add_Header__(cls, doc: DocumentObject, header_text: str, header_order: int) -> None:
        # 处理页眉
        section = doc.sections[header_order - 1]
        header = section.header
        header.is_linked_to_previous = False  # 断开链接
        
        # 删除原有内容
        for paragraph in header.paragraphs:
            p_element = paragraph._element
            p_element.getparent().remove(p_element)
        header_para = header.add_paragraph()

        # 添加页眉横线
        pPr = header_para._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')   # 实线
        bottom.set(qn('w:sz'), '6')         # 6 * 1/8 = 0.75 磅（Word 默认值）
        bottom.set(qn('w:space'), '1')      # 与文字间距 1 磅
        bottom.set(qn('w:color'), 'auto')   # 自动颜色（黑）
        pBdr.append(bottom)
        pPr.append(pBdr)

        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_text = f"{cls.software_name} {cls.software_version} 源代码 第{header_order}部分 {header_text}"
        header_run = header_para.add_run(header_text)
        
        # 设置字体样式
        header_run.font.size = Pt(9)
        header_run.font.name = "Times New Roman"
        header_run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
    

    # 添加代码
    @classmethod
    def __Add_Code__(cls, doc: DocumentObject, code_text: str) -> None:
        code_para = doc.add_paragraph()
        code_para_format = code_para.paragraph_format
        code_para_format.line_spacing = Pt(15)
        code_run = code_para.add_run(code_text + f"\n\n        *********END OF PART {cls.current_codeblock}*********")
        code_run.font.name = "Consolas"
        code_run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
        code_run.font.size = Pt(10.5)
        
        if cls.current_codeblock != cls.total_codeblock:  # 最后一个代码块
            doc.add_section(WD_SECTION.NEW_PAGE)
    

    # [Ci] 添加代码块（一整个代码块，一步搞定标题、页眉、代码，并分页）
    @classmethod
    def Add_CodeBlock(cls, doc: DocumentObject, code_text: str, code_name: str) -> DocumentObject:
        cls.current_codeblock += 1
        DocumentWriter.__Add_Title__(doc, code_name, cls.current_codeblock)
        DocumentWriter.__Add_Header__(doc, code_name, cls.current_codeblock)
        DocumentWriter.__Add_Code__(doc, code_text)
        
        return doc


    def Save(doc: DocumentObject, path: str) -> None:
        doc.save(path)


    @classmethod
    def Generate_One(cls, doc: DocumentObject, code_path: list) -> tuple[bool|str]:
        code_name = File_Manager.Get_FileName(code_path)
        
        if code_name == None:
            print("[ERROR] Cannot detect Code Name, please use absolute path instead of relative path")
            return False, "文件名获取失败"
        
        code_text = File_Manager.Read_File(code_path)
        if code_text == None: 
            print("[Waring] No text in this file")
            return False, "文件内容读取失败，或空文件！"

        DocumentWriter.Add_CodeBlock(doc, code_text, code_name)
        print(f"[INFO] Add Code Block {cls.current_codeblock}/{cls.total_codeblock}")
        return True, f"添加成功{cls.current_codeblock}/{cls.total_codeblock}"

    

    
if __name__ == "__main__":
    # Unit Test
    DocumentWriter.total_codeblock = 2
    doc = DocumentWriter.Create()

    DocumentWriter.Generate_One(doc, "chores/main2.py")
    DocumentWriter.Generate_One(doc, "chores/main3.py")

    DocumentWriter.Save(doc, "chores/test.docx")
